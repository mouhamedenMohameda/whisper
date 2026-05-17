import io
import re
import logging
from datetime import datetime
import asyncio
from typing import Annotated, Optional

from credits_wallet import debit_credits
from database import get_db
from deps import require_wallet_user
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session

from models import User
from pricing import (
    billed_mru_to_wallet_units_debit,
    export_job_billed,
    export_premium_billed,
    wallet_units_to_mru_display,
    latex_conversion_billed_mru,
)
from course_from_transcript import _groq_client, _groq_chat
from groq import Groq
import os
import subprocess
import tempfile
import shutil

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    HRFlowable,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

logger = logging.getLogger(__name__)

router = APIRouter()


class ExportRequest(BaseModel):
    lesson: str
    subject: str = "Lecture"
    filename: str = "lesson"
    language: str = "fr"  # 'fr' or 'ar'


def strip_markdown(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"`(.*?)`", r"\1", text)
    text = re.sub(r"#{1,6}\s+", "", text)
    return text.strip()


@router.post("/export/pdf")
async def export_pdf(
    req: ExportRequest,
    db: Annotated[Session, Depends(get_db)],
    _u: Annotated[Optional[User], Depends(require_wallet_user)],
):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=2 * cm,
        leftMargin=2 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2 * cm,
        title=req.subject,
    )

    styles = getSampleStyleSheet()
    primary = colors.HexColor("#1a1a2e")
    accent = colors.HexColor("#4f46e5")
    light = colors.HexColor("#f0f0ff")
    muted = colors.HexColor("#6b7280")

    title_style = ParagraphStyle(
        "LTitle",
        fontSize=26,
        textColor=primary,
        spaceAfter=6,
        leading=32,
        fontName="Helvetica-Bold",
        alignment=TA_CENTER,
    )
    subtitle_style = ParagraphStyle(
        "LSub",
        fontSize=12,
        textColor=muted,
        spaceAfter=4,
        alignment=TA_CENTER,
    )
    h2_style = ParagraphStyle(
        "LH2",
        fontSize=15,
        textColor=accent,
        spaceBefore=18,
        spaceAfter=8,
        fontName="Helvetica-Bold",
    )
    h3_style = ParagraphStyle(
        "LH3",
        fontSize=12,
        textColor=primary,
        spaceBefore=12,
        spaceAfter=6,
        fontName="Helvetica-Bold",
    )
    body_style = ParagraphStyle(
        "LBody",
        fontSize=10,
        textColor=primary,
        leading=16,
        spaceAfter=6,
        alignment=TA_LEFT,
    )
    bold_style = ParagraphStyle(
        "LBold",
        fontSize=10,
        textColor=primary,
        leading=16,
        spaceAfter=4,
        fontName="Helvetica-Bold",
    )
    bullet_style = ParagraphStyle(
        "LBullet",
        fontSize=10,
        textColor=primary,
        leading=16,
        spaceAfter=4,
        leftIndent=16,
        bulletIndent=4,
    )

    story: list = []

    story.append(Spacer(1, 3 * cm))
    story.append(HRFlowable(width="100%", thickness=3, color=accent))
    story.append(Spacer(1, 0.5 * cm))
    story.append(
        Paragraph(
            "LecturAI",
            ParagraphStyle(
                "Brand",
                fontSize=11,
                textColor=accent,
                alignment=TA_CENTER,
                fontName="Helvetica-Bold",
            ),
        )
    )
    story.append(Spacer(1, 0.3 * cm))

    is_ar = req.language.lower().startswith("ar")
    label_subject = "المادة" if is_ar else "Sujet"
    label_gen = "تاريخ التوليد" if is_ar else "Généré le"
    
    # Titre dynamique (cherche le début de la section 1, peu importe la langue)
    title_match = re.search(r"##\s*1\.\s*.*?\n\s*(.+)", req.lesson, re.MULTILINE)
    lesson_title = title_match.group(1).strip() if title_match else req.subject

    story.append(Paragraph(lesson_title.replace("&", "&amp;"), title_style))
    story.append(Spacer(1, 0.3 * cm))
    story.append(Paragraph(f"{label_subject}: {req.subject}", subtitle_style))
    story.append(
        Paragraph(
            f"{label_gen}: {datetime.now().strftime('%d/%m/%Y')}",
            subtitle_style,
        )
    )
    story.append(Spacer(1, 0.5 * cm))
    story.append(HRFlowable(width="100%", thickness=1, color=light))
    story.append(PageBreak())

    lines = req.lesson.split("\n")
    i = 0
    in_table = False
    table_rows: list[list[str]] = []

    while i < len(lines):
        line = lines[i]

        if line.startswith("## "):
            if in_table and table_rows:
                _flush_table_pdf(story, table_rows, accent, light)
                in_table = False
                table_rows = []
            story.append(Paragraph(strip_markdown(line[3:]), h2_style))

        elif line.startswith("# ") and not line.startswith("## "):
            if in_table and table_rows:
                _flush_table_pdf(story, table_rows, accent, light)
                in_table = False
                table_rows = []
            # H1 : même style que H2 mais un peu plus grand
            h1_style = ParagraphStyle(
                "LH1",
                fontSize=18,
                textColor=accent,
                spaceBefore=22,
                spaceAfter=10,
                fontName="Helvetica-Bold",
            )
            story.append(Paragraph(strip_markdown(line[2:]), h1_style))

        elif line.startswith("### "):
            if in_table and table_rows:
                _flush_table_pdf(story, table_rows, accent, light)
                in_table = False
                table_rows = []
            story.append(Paragraph(strip_markdown(line[4:]), h3_style))

        elif line.startswith("| ") and "|" in line:
            if not in_table:
                in_table = True
                table_rows = []
            if "---" not in line:
                cells = [strip_markdown(c.strip()) for c in line.split("|") if c.strip()]
                table_rows.append(cells)
        else:
            if in_table and table_rows:
                _flush_table_pdf(story, table_rows, accent, light)
                in_table = False
                table_rows = []

            if line.startswith("- ") or line.startswith("* "):
                text = strip_markdown(line[2:])
                story.append(Paragraph(f"• {text}", bullet_style))
            elif re.match(r"^\*\*(.+?)\*\*\s*[—\-]", line):
                story.append(Paragraph(strip_markdown(line), body_style))
            elif line.strip() == "---":
                story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#e0e0e0")))
                story.append(Spacer(1, 0.2 * cm))
            elif line.strip():
                story.append(Paragraph(strip_markdown(line), bold_style if line.startswith("**") else body_style))
            else:
                story.append(Spacer(1, 0.15 * cm))

        i += 1

    if in_table and table_rows:
        _flush_table_pdf(story, table_rows, accent, light)

    await asyncio.to_thread(doc.build, story)
    buffer.seek(0)

    _, billed_mru = export_job_billed()
    charged = billed_mru_to_wallet_units_debit(billed_mru)
    if charged > 0:
        debit_credits(db, _u, charged)

    safe_name = re.sub(r"[^\w\-]", "_", req.filename)
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{safe_name}_lesson.pdf"'},
    )


def _flush_table_pdf(story, table_rows, accent, light):
    if len(table_rows) <= 1:
        return
    col_count = max(len(r) for r in table_rows)
    padded = [r + [""] * (col_count - len(r)) for r in table_rows]
    col_width = (A4[0] - 4 * cm) / col_count
    t = Table(padded, colWidths=[col_width] * col_count)
    t.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), accent),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, light]),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#e0e0e0")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
            ]
        )
    )
    story.append(t)
    story.append(Spacer(1, 0.3 * cm))


@router.post("/export/docx")
async def export_docx(
    req: ExportRequest,
    db: Annotated[Session, Depends(get_db)],
    _u: Annotated[Optional[User], Depends(require_wallet_user)],
):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover.add_run("LecturAI")
    run.font.color.rgb = RGBColor(79, 70, 229)
    run.font.size = Pt(11)
    run.font.bold = True

    is_ar = req.language.lower().startswith("ar")
    label_subject = "المادة" if is_ar else "Sujet"
    label_gen = "تاريخ التوليد" if is_ar else "Généré le"

    title_match = re.search(r"##\s*1\.\s*.*?\n\s*(.+)", req.lesson, re.MULTILINE)
    lesson_title = title_match.group(1).strip() if title_match else req.subject

    h = doc.add_heading(lesson_title, 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph(
        f"{label_subject}: {req.subject}  |  {label_gen}: {datetime.now().strftime('%d/%m/%Y')}"
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.color.rgb = RGBColor(107, 114, 128)

    doc.add_page_break()

    table_buffer: list[list[str]] = []
    collecting_table = False

    def flush_table():
        nonlocal collecting_table, table_buffer
        if not table_buffer:
            collecting_table = False
            table_buffer = []
            return
        rows = len(table_buffer)
        cols = max(len(r) for r in table_buffer)
        table = doc.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"
        for ri, row in enumerate(table_buffer):
            for ci in range(cols):
                cell_text = row[ci] if ci < len(row) else ""
                table.rows[ri].cells[ci].text = cell_text
        doc.add_paragraph("")
        collecting_table = False
        table_buffer = []

    for line in req.lesson.split("\n"):
        if line.startswith("## "):
            flush_table()
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("# ") and not line.startswith("## "):
            flush_table()
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith("### "):
            flush_table()
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("| ") and "|" in line:
            if "---" in line:
                continue
            cells = [strip_markdown(c.strip()) for c in line.split("|") if c.strip()]
            if cells:
                collecting_table = True
                table_buffer.append(cells)
        elif line.startswith("- ") or line.startswith("* "):
            flush_table()
            doc.add_paragraph(strip_markdown(line[2:]), style="List Bullet")
        elif line.strip() == "---":
            flush_table()
            doc.add_paragraph("─" * 60)
        elif line.strip():
            flush_table()
            doc.add_paragraph(strip_markdown(line))
        else:
            flush_table()

    flush_table()

    buffer = io.BytesIO()
    await asyncio.to_thread(doc.save, buffer)
    buffer.seek(0)

    _, billed_mru_d = export_job_billed()
    charged_d = billed_mru_to_wallet_units_debit(billed_mru_d)
    if charged_d > 0:
        debit_credits(db, _u, charged_d)

    safe_name = re.sub(r"[^\w\-]", "_", req.filename)
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{safe_name}_lesson.docx"'},
    )


LATEX_PREMIUM_SYSTEM_PROMPT = """You are an Elite Academic Publisher and master LaTeX Typographer.
The user is paying for an ULTRA-PREMIUM experience. The document must be visually stunning, perfectly balanced, and evoke the quality of an expensive textbook from a top-tier university (Harvard/Oxford style).

# Visual & Typographic Requirements:
- Use 'article' class with [11pt, a4paper].
- Essential Packages:
    - 'geometry' (left=2cm, right=2cm, top=2.5cm, bottom=2.5cm).
    - 'libertine', 'biolinum', 'inconsolata'.
    - 'lettrine', 'microtype', 'xcolor' (dvipsnames), 'titlesec', 'fancyhdr', 'booktabs', 'multicol', 'enumitem', 'tabularx'.
    - 'tcolorbox' (v4+) with 'skins' and 'breakable'.

# Custom Elite Styles (MUST use these definitions in the preamble):
1. \definecolor{primary}{HTML}{1e1b4b} \definecolor{accent}{HTML}{4f46e5} \definecolor{bglight}{HTML}{f8fafc}
2. \newtcolorbox{DefinitionBox}[1]{colback=accent!5, colframe=accent, fonttitle=\bfseries\sffamily, title=#1, arc=4pt, breakable, enhanced, shadow={1mm}{-1mm}{0mm}{accent!20}}
3. \newtcolorbox{TakeawayBox}{colback=bglight, colframe=accent, leftrule=4pt, rightrule=0pt, toprule=0pt, bottomrule=0pt, sharp corners, breakable}
4. \newtcolorbox{WarningBox}[1]{colback=red!5, colframe=red!75!black, fonttitle=\bfseries\sffamily, title=#1, arc=4pt, breakable}

# Layout Rules:
- COVER PAGE: Use \begin{titlepage}, \centering, \vspace*{4cm}, Massive Title in \Huge\bfseries\color{primary}, \vfill, "Édition Premium par LecturAI" in \small\scshape, \end{titlepage}.
- TABLES: ALWAYS use 'tabularx' with width '\linewidth' and at least one 'X' column for long text to prevent margin overflow. Use 'booktabs' rules (\toprule, \midrule, \bottomrule).
- HEADERS: Fancyhdr with 'LecturAI Elite Series' and section names.
- SECTIONING: Large Biolinum Bold fonts with a thin rule below.
- ENUMERATION: Use enumitem for tight, professional lists.
- TYPOGRAPHY: Use \lettrine for the first letter of the introduction.

# Technical Rigor:
- ESCAPE ALL SPECIAL CHARACTERS: (%, &, $, #, _, {, }, ~, ^, \\). 
- If text contains equations, use AMS-LaTeX.
- Ensure perfect French support via babel.
- DANGER: DO NOT WRAP THE OUTPUT IN MARKDOWN CODE BLOCKS (```latex). 
- OUTPUT ONLY THE RAW LATEX SOURCE CODE starting with \documentclass. No backticks.
"""



@router.post("/export/pdf-professional/estimate")
async def estimate_premium_pdf(
    req: ExportRequest,
    _u: Annotated[Optional[User], Depends(require_wallet_user)],
):
    """Estime le coût du PDF Premium basé sur un multiplicateur x8 du coût IA."""
    _, billed_mru = latex_conversion_billed_mru(req.lesson)
    return {
        "cost_mru": billed_mru,
        "cost_display": f"{wallet_units_to_mru_display(billed_mru_to_wallet_units_debit(billed_mru))} MRU",
        "has_credits": _u.credit_balance >= billed_mru_to_wallet_units_debit(billed_mru) if _u else False,
    }


@router.post("/export/pdf-professional/generate")
async def export_pdf_professional(
    req: ExportRequest,
    db: Annotated[Session, Depends(get_db)],
    _u: Annotated[Optional[User], Depends(require_wallet_user)],
):
    """Génère un PDF professionnel via LaTeX avec facturation x8 du coût réel.

    Pipeline LaTeX déporté dans ``pdf_premium.generate_premium_pdf_bytes`` (module partagé
    avec le bot WhatsApp). Cette route conserve la logique HTTP + facturation + streaming.
    """
    from pdf_premium import (
        generate_premium_pdf_bytes,
        PremiumPdfError,
        PremiumPdfUnavailable,
    )

    _, billed_mru = latex_conversion_billed_mru(req.lesson)
    charged = billed_mru_to_wallet_units_debit(billed_mru)

    if _u and _u.credit_balance < charged:
        raise HTTPException(status_code=402, detail="Crédits insuffisants pour l'export professionnel.")

    api_key = (os.getenv("GROQ_API_KEY") or "").strip()
    if not api_key:
        raise HTTPException(status_code=500, detail="Clé API manquante pour la conversion LaTeX.")

    try:
        pdf_data, _page_count, _latex = await asyncio.to_thread(
            generate_premium_pdf_bytes,
            lesson_markdown=req.lesson,
            subject=req.subject,
            language=req.language,
            api_key=api_key,
        )
    except PremiumPdfUnavailable as exc:
        raise HTTPException(status_code=501, detail=str(exc)) from None
    except PremiumPdfError as exc:
        raise HTTPException(status_code=502, detail=str(exc)) from None

    if charged > 0:
        debit_credits(db, _u, charged)

    buffer = io.BytesIO(pdf_data)
    safe_name = re.sub(r"[^\w\-]", "_", req.filename)
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{safe_name}_professional.pdf"'},
    )

